"""
╔══════════════════════════════════════════════════════════════════╗
║           AI-LLM Notebook — Chatbot Backend (Ollama)            ║
║  FastAPI server with file upload, summarize & chat endpoints     ║
╚══════════════════════════════════════════════════════════════════╝

SETUP:
  pip install fastapi uvicorn python-multipart pymupdf python-docx pillow pytesseract requests

OLLAMA SETUP:
  1. Install Ollama: https://ollama.ai
  2. Pull a model:   ollama pull llama3.2       (ya koi bhi model)
  3. Ollama start karo (usually auto-starts)

RUN:
  python chatbot.py
  → Server: http://127.0.0.1:8000
  → Docs:   http://127.0.0.1:8000/docs
"""

import os
import io
import re
import json
import base64
import requests
import tempfile
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel


# ──────────────────────────────────────────────────────────────────
#  CONFIG — ye values apne hisaab se badlo
# ──────────────────────────────────────────────────────────────────
OLLAMA_BASE_URL  = "http://localhost:11434"   # Ollama ka default URL
OLLAMA_MODEL     = "llama3.2"                 # ollama list se model naam dekho
OLLAMA_TIMEOUT   = 120                        # seconds (bade documents ke liye badao)

# Supported file types
SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "image/webp": "image",
}


# ──────────────────────────────────────────────────────────────────
#  FASTAPI APP
# ──────────────────────────────────────────────────────────────────
app = FastAPI(
    title="AI-LLM Notebook Chatbot API",
    description="Ollama-powered chatbot with document upload & summarization",
    version="1.0.0",
)

# CORS — frontend ko allow karo
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],       # Production mein apna domain dalo
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────────────────────────────────────────────────────────
#  PYDANTIC MODELS
# ──────────────────────────────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = "No document context. Answer generally."

class ChatResponse(BaseModel):
    reply: str
    model: str

class SummarizeResponse(BaseModel):
    summary: str
    title: str
    file_name: str
    file_type: str
    char_count: int


# ──────────────────────────────────────────────────────────────────
#  OLLAMA HELPER
# ──────────────────────────────────────────────────────────────────
def ollama_generate(prompt: str, system: str = "") -> str:
    """
    Ollama /api/generate endpoint call karta hai.
    Streaming OFF rakha hai taaki ek saath response mile.
    """
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.7,
            "top_p": 0.9,
            "num_predict": 1024,
        }
    }
    if system:
        payload["system"] = system

    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json=payload,
            timeout=OLLAMA_TIMEOUT
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("response", "").strip()
    except requests.exceptions.ConnectionError:
        raise HTTPException(
            status_code=503,
            detail="❌ Ollama connect nahi ho raha. 'ollama serve' command run karo terminal mein."
        )
    except requests.exceptions.Timeout:
        raise HTTPException(
            status_code=504,
            detail="❌ Ollama timeout ho gaya. Model load ho raha hoga, thoda wait karo."
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ollama error: {str(e)}")


def check_ollama_health() -> dict:
    """Ollama running hai ya nahi check karta hai."""
    try:
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)
        models = [m["name"] for m in resp.json().get("models", [])]
        return {"status": "ok", "models": models}
    except Exception:
        return {"status": "offline", "models": []}


# ──────────────────────────────────────────────────────────────────
#  TEXT EXTRACTION FUNCTIONS
# ──────────────────────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """PDF se text nikalta hai — PyMuPDF (fitz) use karta hai."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")
        doc.close()
        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            return "⚠️ PDF mein readable text nahi mila (scanned image ho sakta hai)."
        return full_text
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF install nahi hai. Run: pip install pymupdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF read error: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """DOCX se text nikalta hai — python-docx use karta hai."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tables bhi extract karo
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx install nahi hai. Run: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX read error: {str(e)}")


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Image se text nikalta hai.
    Ollama vision model use karta hai agar available ho,
    warna pytesseract fallback.
    """
    # Try Ollama vision model (llava etc.)
    health = check_ollama_health()
    vision_models = [m for m in health.get("models", []) if any(
        v in m.lower() for v in ["llava", "bakllava", "moondream", "minicpm"]
    )]

    if vision_models:
        try:
            img_b64 = base64.b64encode(file_bytes).decode("utf-8")
            payload = {
                "model": vision_models[0],
                "prompt": "Extract and describe all text visible in this image. List everything you can read.",
                "images": [img_b64],
                "stream": False,
            }
            resp = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json=payload, timeout=60)
            resp.raise_for_status()
            return resp.json().get("response", "").strip()
        except Exception:
            pass  # Fallback to pytesseract

    # Pytesseract fallback
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img, lang="eng+hin")  # Hindi + English
        return text.strip() if text.strip() else "⚠️ Image mein readable text nahi mila."
    except ImportError:
        return "⚠️ Image text extraction ke liye pytesseract ya llava model chahiye."
    except Exception as e:
        return f"⚠️ Image read error: {str(e)}"


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> tuple[str, str]:
    """
    File type ke hisaab se text extract karta hai.
    Returns: (extracted_text, file_type_label)
    """
    ct = content_type.lower()
    fname = filename.lower()

    # Content-type se detect karo
    if ct == "application/pdf" or fname.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes), "pdf"
    elif ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or fname.endswith(".docx"):
        return extract_text_from_docx(file_bytes), "docx"
    elif ct.startswith("image/") or fname.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return extract_text_from_image(file_bytes), "image"
    elif ct == "text/plain" or fname.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace"), "txt"
    else:
        # Try as plain text
        try:
            return file_bytes.decode("utf-8", errors="replace"), "txt"
        except Exception:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {content_type}")


def truncate_text(text: str, max_chars: int = 8000) -> str:
    """LLM context limit ke liye text truncate karta hai."""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return text[:half] + f"\n\n... [Document truncated — {len(text) - max_chars} chars removed] ...\n\n" + text[-half:]


# ──────────────────────────────────────────────────────────────────
#  API ENDPOINTS
# ──────────────────────────────────────────────────────────────────

@app.get("/", tags=["Health"])
def root():
    """Server health check."""
    ollama = check_ollama_health()
    return {
        "status": "✅ Server running",
        "model": OLLAMA_MODEL,
        "ollama": ollama["status"],
        "available_models": ollama.get("models", []),
        "endpoints": ["/chat", "/summarize", "/health", "/models"]
    }


@app.get("/health", tags=["Health"])
def health():
    """Detailed health check."""
    ollama = check_ollama_health()
    return {
        "server": "ok",
        "ollama_status": ollama["status"],
        "ollama_url": OLLAMA_BASE_URL,
        "active_model": OLLAMA_MODEL,
        "available_models": ollama.get("models", []),
    }


@app.get("/models", tags=["Models"])
def list_models():
    """Available Ollama models list karo."""
    health = check_ollama_health()
    if health["status"] == "offline":
        raise HTTPException(status_code=503, detail="Ollama offline hai.")
    return {"models": health["models"], "current": OLLAMA_MODEL}


@app.post("/summarize", response_model=SummarizeResponse, tags=["Documents"])
async def summarize_document(file: UploadFile = File(...)):
    """
    Document upload karo → text extract karo → Ollama se summary banao.
    
    Supported: PDF, DOCX, TXT, PNG, JPG, JPEG, WEBP
    """
    # File bytes read karo
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty file upload kiya hai.")

    # File size check (50MB max)
    max_size = 50 * 1024 * 1024
    if len(file_bytes) > max_size:
        raise HTTPException(status_code=413, detail="File too large. Max 50MB allowed.")

    # Text extract karo
    content_type = file.content_type or "application/octet-stream"
    raw_text, file_type = extract_text(file_bytes, content_type, file.filename or "file")

    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="File mein koi readable text nahi mila.")

    # Truncate karo LLM ke liye
    truncated = truncate_text(raw_text, max_chars=6000)

    # Summary generate karo
    system_prompt = """Tu ek expert document analyzer hai. 
Tujhe documents ki clear, structured summary banani hai Hindi aur English mix mein (Hinglish).
Summary mein ye cheezein zaroor shamil kar:
- Main points (bullet points mein)
- Key findings ya conclusions  
- Important numbers ya facts
- Overall context

Format: Clean aur readable. Markdown use kar."""

    user_prompt = f"""Ye document ka text hai:

---
{truncated}
---

Is document ki ek comprehensive summary banao. 
Pehle ek short title do (1 line), phir detailed summary."""

    summary = ollama_generate(user_prompt, system_prompt)

    # Title extract karo summary se
    lines = summary.strip().split("\n")
    title = file.filename or "Document"
    for line in lines[:3]:
        clean = line.strip().lstrip("#").strip().lstrip("*").strip()
        if clean and len(clean) < 100 and not clean.endswith(":"):
            title = clean
            break

    return SummarizeResponse(
        summary=summary,
        title=title,
        file_name=file.filename or "uploaded_file",
        file_type=file_type,
        char_count=len(raw_text),
    )


@app.post("/chat", response_model=ChatResponse, tags=["Chat"])
async def chat(request: ChatRequest):
    """
    User message + document context ke saath chat karo.
    
    Body:
        message: User ka question
        context: Document summaries (frontend se automatically aata hai)
    """
    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message empty nahi hona chahiye.")

    # System prompt
    system_prompt = """Tu ek helpful AI assistant hai jo document questions answer karta hai.

Rules:
1. Context mein diye gaye documents ke basis pe jawab do
2. Agar context mein answer nahi hai toh clearly bolo
3. Hinglish mein jawab do (Hindi + English mix) — friendly tone
4. Accurate aur helpful raho
5. Agar user document ke baare mein pooche jo context mein nahi hai, toh clearly indicate karo
6. Technical terms English mein rakh sakte ho
7. Structured jawab do — bullet points, headings use karo jab zaroorat ho"""

    # Full prompt with context
    user_prompt = f"""=== DOCUMENT CONTEXT ===
{request.context}

=== USER QUESTION ===
{request.message}

=== INSTRUCTION ===
Context ke basis pe helpful jawab do. Agar context mein relevant information hai toh use karo.
Agar nahi hai toh general knowledge se jawab do aur clearly batao ki ye document-based nahi hai."""

    reply = ollama_generate(user_prompt, system_prompt)

    return ChatResponse(reply=reply, model=OLLAMA_MODEL)


# ──────────────────────────────────────────────────────────────────
#  EXTRA UTILITY ENDPOINTS
# ──────────────────────────────────────────────────────────────────

@app.post("/extract-text", tags=["Documents"])
async def extract_text_only(file: UploadFile = File(...)):
    """
    Sirf text extract karo — summary nahi.
    Useful for debugging ya raw text dekhne ke liye.
    """
    file_bytes = await file.read()
    content_type = file.content_type or "text/plain"
    raw_text, file_type = extract_text(file_bytes, content_type, file.filename or "file")
    return {
        "text": raw_text[:5000],  # First 5000 chars
        "total_chars": len(raw_text),
        "file_type": file_type,
        "truncated": len(raw_text) > 5000,
    }


@app.post("/chat-stream", tags=["Chat"])
async def chat_stream(request: ChatRequest):
    """
    Streaming chat endpoint (experimental).
    Frontend mein EventSource ya fetch streaming ke liye.
    """
    from fastapi.responses import StreamingResponse

    system_prompt = """Tu ek helpful AI assistant hai jo document questions answer karta hai.
Hinglish mein jawab do. Context ke basis pe accurate jawab do."""

    user_prompt = f"""Context:\n{request.context}\n\nQuestion: {request.message}"""

    def generate():
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": user_prompt,
            "system": system_prompt,
            "stream": True,
            "options": {"temperature": 0.7}
        }
        try:
            with requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json=payload,
                stream=True,
                timeout=OLLAMA_TIMEOUT
            ) as resp:
                for line in resp.iter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            token = data.get("response", "")
                            if token:
                                yield f"data: {json.dumps({'token': token})}\n\n"
                            if data.get("done"):
                                yield "data: [DONE]\n\n"
                                break
                        except json.JSONDecodeError:
                            continue
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ──────────────────────────────────────────────────────────────────
#  STARTUP EVENT
# ──────────────────────────────────────────────────────────────────
@app.on_event("startup")
async def startup_event():
    print("\n" + "═" * 60)
    print("  🤖 AI-LLM Notebook — Chatbot Backend")
    print("═" * 60)
    print(f"  Model    : {OLLAMA_MODEL}")
    print(f"  Ollama   : {OLLAMA_BASE_URL}")

    health = check_ollama_health()
    if health["status"] == "ok":
        print(f"  Ollama   : ✅ Online")
        print(f"  Models   : {', '.join(health['models']) or 'None pulled yet'}")
        if OLLAMA_MODEL not in " ".join(health["models"]):
            print(f"\n  ⚠️  WARNING: '{OLLAMA_MODEL}' model nahi mila!")
            print(f"  Run: ollama pull {OLLAMA_MODEL}")
    else:
        print(f"  Ollama   : ❌ Offline — 'ollama serve' run karo")

    print("\n  Endpoints:")
    print("  → POST /chat          Chat with context")
    print("  → POST /summarize     Upload & summarize document")
    print("  → GET  /health        Health check")
    print("  → GET  /docs          Swagger UI")
    print("═" * 60 + "\n")


# ──────────────────────────────────────────────────────────────────
#  MAIN — Direct run karne ke liye
# ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "chatbot:app",
        host="127.0.0.1",
        port=8000,
        reload=True,        # Development mein auto-reload
        log_level="info",
    )